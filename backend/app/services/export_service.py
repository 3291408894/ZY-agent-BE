"""
通用导出服务 — Word/PDF 导出 (功能2 & 功能1 共用)
"""

import os
import re
from datetime import datetime

from app.core.config import settings
from app.schemas.exam_paper import ExportFormat


def _clean_math(text: str) -> str:
    """去除 LaTeX 数学标记，转为可打印的纯文本"""
    # 移除 $ 和 $$ 标记，保留内部内容
    text = re.sub(r'\$\$(.*?)\$\$', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    # 替换常见 LaTeX 命令为 Unicode
    text = text.replace('\\frac', '')
    text = text.replace('\\dfrac', '')
    text = text.replace('\\sqrt', '√')
    text = text.replace('\\log', 'log')
    text = text.replace('\\ln', 'ln')
    text = text.replace('\\sin', 'sin')
    text = text.replace('\\cos', 'cos')
    text = text.replace('\\tan', 'tan')
    text = text.replace('\\left', '')
    text = text.replace('\\right', '')
    text = text.replace('\\cdot', '·')
    text = text.replace('\\times', '×')
    text = text.replace('\\div', '÷')
    text = text.replace('\\pm', '±')
    text = text.replace('\\leq', '≤')
    text = text.replace('\\geq', '≥')
    text = text.replace('\\neq', '≠')
    text = text.replace('\\infty', '∞')
    text = text.replace('\\alpha', 'α')
    text = text.replace('\\beta', 'β')
    text = text.replace('\\pi', 'π')
    text = text.replace('\\theta', 'θ')
    text = text.replace('\\degree', '°')
    text = re.sub(r'[{}]', '', text)
    text = re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    # 清理多余空格
    text = re.sub(r'\s+', ' ', text).strip()
    return text


class ExportService:
    """通用导出服务，支持试卷和教案的 Word/PDF 导出"""

    def __init__(self):
        self.export_dir = os.path.join(settings.UPLOAD_DIR, "exports")
        os.makedirs(self.export_dir, exist_ok=True)

    def export_exam_paper(
        self,
        paper_data,
        export_format: ExportFormat,
    ) -> tuple[str, str]:
        """
        导出试卷。

        Returns:
            tuple[str, str]: (文件路径, 文件名)
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
        except ImportError:
            raise ImportError("请安装 python-docx 库: pip install python-docx")

        doc = Document()

        # 设置 A4 纸张
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # 构建试卷内容
        content = paper_data.content or {}

        # ── 密封线 ──
        self._add_seal_line(doc, paper_data)

        # ── 试卷头部 ──
        header = content.get("header", {})
        self._add_header(doc, header)

        # ── 考生须知 ──
        instructions = header.get("instructions", "")
        if instructions:
            self._add_instructions(doc, instructions)

        # ── 试题部分 ──
        sections = content.get("sections", [])
        for sec in sections:
            self._add_section(doc, sec)

        # ── 分页后添加参考答案 ──
        doc.add_page_break()
        self._add_answer_section(doc, content)

        # ── 保存文件 ──
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = paper_data.title.replace(" ", "_").replace("/", "_")[:50]
        file_name = f"{safe_title}_{timestamp}.docx"
        file_path = os.path.join(self.export_dir, file_name)
        doc.save(file_path)

        return file_path, file_name

    # ─────────────────────────────────────────────
    # 密封线
    # ─────────────────────────────────────────────

    def _add_seal_line(self, doc, paper_data):
        """添加密封线"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("┌────────────────── 密封线内不要答题 ──────────────────┐")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lines = [
            f"学校：_______________    班级：_______________",
            f"姓名：_______________    学号：_______________",
        ]
        for line in lines:
            run2 = p2.add_run(line + "\n")
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(128, 128, 128)

        p3 = doc.add_paragraph()
        run3 = p3.add_run("└──────────────────────────────────────────────────────┘")
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(128, 128, 128)

        # 分隔线
        doc.add_paragraph("─" * 60)

    # ─────────────────────────────────────────────
    # 试卷头部
    # ─────────────────────────────────────────────

    def _add_header(self, doc, header: dict):
        """添加试卷标题区域"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 标题
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(header.get("title", ""))
        title_run.bold = True
        title_run.font.size = Pt(18)

        # 信息行
        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_lines = [
            f"学科：{header.get('subject', '')}　　"
            f"年级：{header.get('grade', '')}　　"
            f"考试类型：{header.get('exam_type', '')}",
            f"总分：{header.get('total_score', '')}分　　"
            f"考试时间：{header.get('duration_minutes', 120)}分钟",
        ]
        for line in info_lines:
            info_run = info_p.add_run(line + "\n")
            info_run.font.size = Pt(11)

    # ─────────────────────────────────────────────
    # 考生须知
    # ─────────────────────────────────────────────

    def _add_instructions(self, doc, instructions: str):
        """添加考生须知"""
        from docx.shared import Pt

        inst_p = doc.add_paragraph()
        inst_run = inst_p.add_run("【考生须知】")
        inst_run.bold = True
        inst_run.font.size = Pt(10)

        content_p = doc.add_paragraph()
        content_run = content_p.add_run(instructions)
        content_run.font.size = Pt(10)

    # ─────────────────────────────────────────────
    # 试题大题
    # ─────────────────────────────────────────────

    def _add_section(self, doc, section_data: dict):
        """添加一个大题（如：一、选择题）"""
        from docx.shared import Pt

        # 大题标题
        sec_p = doc.add_paragraph()
        sec_run = sec_p.add_run(section_data.get("title", ""))
        sec_run.bold = True
        sec_run.font.size = Pt(14)

        # 答题说明
        instructions = section_data.get("instructions", "")
        if instructions:
            inst_p = doc.add_paragraph()
            inst_run = inst_p.add_run(instructions)
            inst_run.font.size = Pt(10)
            inst_run.italic = True

        # 题目列表
        questions = section_data.get("questions", [])
        for q in questions:
            self._add_question(doc, q)

    # ─────────────────────────────────────────────
    # 单道试题
    # ─────────────────────────────────────────────

    def _add_question(self, doc, question: dict):
        """添加一道试题"""
        from docx.shared import Pt

        number = question.get("number", "")
        stem = question.get("stem", "")
        score = question.get("score", 0)
        q_type = question.get("question_type", "")
        options = question.get("options")

        # 题号 + 题目
        q_p = doc.add_paragraph()
        q_text = f"{number}. {_clean_math(stem)}"
        if score:
            q_text += f" （{score}分）"
        q_run = q_p.add_run(q_text)
        q_run.font.size = Pt(12)

        # 选项（选择题）
        if options and q_type == "choice":
            opt_p = doc.add_paragraph()
            for i, opt in enumerate(options):
                indent = "    " if i % 2 == 0 else "        "
                opt_run = opt_p.add_run(f"{indent}{_clean_math(opt)}")
                opt_run.font.size = Pt(12)
                if i % 2 == 1:
                    opt_p.add_run("\n")

        # 作答区留白
        if q_type in ("fill", "fill_blank"):
            blank_p = doc.add_paragraph()
            blank_run = blank_p.add_run("答：___________________________")
            blank_run.font.size = Pt(12)

        elif q_type in ("short_answer", "calculation"):
            # 预留5行作答区
            for _ in range(5):
                space_p = doc.add_paragraph()
                space_p.add_run("_" * 80).font.size = Pt(8)

        elif q_type in ("comprehensive", "analysis"):
            # 预留10行作答区
            for _ in range(10):
                space_p = doc.add_paragraph()
                space_p.add_run("_" * 80).font.size = Pt(8)

    # ─────────────────────────────────────────────
    # 参考答案
    # ─────────────────────────────────────────────

    def _add_answer_section(self, doc, content: dict):
        """添加参考答案和评分标准"""
        from docx.shared import Pt

        # 参考答案标题
        ans_title = doc.add_paragraph()
        ans_run = ans_title.add_run("参考答案与评分标准")
        ans_run.bold = True
        ans_run.font.size = Pt(16)
        ans_title.alignment = 1  # center

        # 逐题答案
        answer_key = content.get("answer_key", [])
        if answer_key:
            for item in answer_key:
                num = item.get("number", "")
                ans = item.get("answer", "")
                sc = item.get("score", "")
                ans_p = doc.add_paragraph()
                ans_text = f"{num}. {ans}"
                if sc:
                    ans_text += f" （{sc}分）"
                ans_run2 = ans_p.add_run(ans_text)
                ans_run2.font.size = Pt(11)

        # 如果没有answer_key，从sections中提取
        if not answer_key:
            sections = content.get("sections", [])
            for sec in sections:
                sec_title = sec.get("title", "")
                sec_p = doc.add_paragraph()
                sec_run = sec_p.add_run(sec_title)
                sec_run.bold = True
                sec_run.font.size = Pt(12)

                for q in sec.get("questions", []):
                    num = q.get("number", "")
                    ans = _clean_math(q.get("answer", ""))
                    sc = q.get("score", "")
                    analysis = q.get("analysis", "")

                    q_p = doc.add_paragraph()
                    q_text = f"{num}. 答案：{ans}"
                    if sc:
                        q_text += f" （{sc}分）"
                    q_run = q_p.add_run(q_text)
                    q_run.font.size = Pt(11)

                    if analysis:
                        a_p = doc.add_paragraph()
                        a_run = a_p.add_run(f"   解析：{_clean_math(analysis)}")
                        a_run.font.size = Pt(10)
                        a_run.italic = True

        # 评分标准
        scoring_guide = content.get("scoring_guide", "")
        if scoring_guide:
            sg_p = doc.add_paragraph()
            sg_run = sg_p.add_run("\n【评分标准】")
            sg_run.bold = True
            sg_run.font.size = Pt(12)

            guide_p = doc.add_paragraph()
            guide_run = guide_p.add_run(scoring_guide)
            guide_run.font.size = Pt(11)
